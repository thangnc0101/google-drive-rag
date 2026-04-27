from __future__ import annotations

import io
import re
from dataclasses import dataclass, field


class UnsupportedFileType(Exception):
    pass


@dataclass
class PageInfo:
    page_number: int
    text: str


@dataclass
class SectionInfo:
    heading: str
    text: str
    level: int = 0


@dataclass
class ParsedDocument:
    filename: str
    text: str
    pages: list[PageInfo] = field(default_factory=list)
    sections: list[SectionInfo] = field(default_factory=list)


class DocumentParser:
    def parse(self, file_bytes: bytes, filename: str, mime_type: str) -> ParsedDocument:
        if mime_type == "application/pdf":
            return self._parse_pdf(file_bytes, filename)
        elif mime_type in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ):
            return self._parse_docx(file_bytes, filename)
        elif mime_type in ("text/plain", "text/markdown"):
            return self._parse_text(file_bytes, filename)
        elif mime_type == "application/vnd.google-apps.document":
            return self._parse_text(file_bytes, filename)
        else:
            raise UnsupportedFileType(f"Unsupported: {mime_type}")

    def _parse_pdf(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        import fitz

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages = []
        all_text_parts = []
        font_sizes: list[tuple[float, str]] = []

        for i, page in enumerate(doc):
            text = page.get_text()
            pages.append(PageInfo(page_number=i + 1, text=text))
            all_text_parts.append(text)
            blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE).get("blocks", [])
            for block in blocks:
                for line in block.get("lines", []):
                    line_text = "".join(span.get("text", "") for span in line.get("spans", []))
                    line_text = line_text.strip()
                    if not line_text:
                        continue
                    sizes = [span["size"] for span in line.get("spans", []) if span.get("text", "").strip()]
                    if sizes:
                        font_sizes.append((max(sizes), line_text))

        doc.close()
        full_text = "\n\n".join(all_text_parts)
        sections = self._extract_pdf_sections(font_sizes, full_text)
        return ParsedDocument(
            filename=filename,
            text=full_text,
            pages=pages,
            sections=sections,
        )

    @staticmethod
    def _extract_pdf_sections(font_sizes: list[tuple[float, str]], full_text: str) -> list[SectionInfo]:
        if not font_sizes:
            return []
        size_counts: dict[float, int] = {}
        for size, text in font_sizes:
            size_counts[size] = size_counts.get(size, 0) + len(text)
        body_size = max(size_counts, key=size_counts.get)

        heading_lines: list[tuple[float, str, int]] = []
        for size, text in font_sizes:
            if size > body_size + 0.5 and len(text) < 200:
                level = 1 if size > body_size + 4 else (2 if size > body_size + 2 else 3)
                heading_lines.append((size, text, level))

        if not heading_lines:
            return []

        sections: list[SectionInfo] = []
        for i, (_, heading, level) in enumerate(heading_lines):
            start_pos = full_text.find(heading)
            if start_pos < 0:
                continue
            content_start = start_pos + len(heading)
            if i + 1 < len(heading_lines):
                next_pos = full_text.find(heading_lines[i + 1][1], content_start)
                body = full_text[content_start:next_pos].strip() if next_pos > content_start else full_text[content_start:].strip()
            else:
                body = full_text[content_start:].strip()
            if body:
                sections.append(SectionInfo(heading=heading, text=body, level=level))

        return sections

    def _parse_docx(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        sections = []
        current_heading = ""
        current_level = 0
        current_texts: list[str] = []
        all_text_parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            all_text_parts.append(text)
            if para.style and para.style.name.startswith("Heading"):
                if current_texts:
                    sections.append(SectionInfo(heading=current_heading, text="\n".join(current_texts), level=current_level))
                current_heading = text
                current_level = self._docx_heading_level(para.style.name)
                current_texts = []
            else:
                current_texts.append(text)

        if current_texts:
            sections.append(SectionInfo(heading=current_heading, text="\n".join(current_texts), level=current_level))

        return ParsedDocument(
            filename=filename,
            text="\n\n".join(all_text_parts),
            sections=sections,
        )

    @staticmethod
    def _docx_heading_level(style_name: str) -> int:
        m = re.search(r"(\d+)$", style_name)
        return int(m.group(1)) if m else 1

    def _parse_text(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        text = file_bytes.decode("utf-8", errors="replace")
        sections = self._extract_markdown_sections(text)
        if not sections:
            sections = self._extract_plaintext_sections(text)
        return ParsedDocument(filename=filename, text=text, sections=sections)

    @staticmethod
    def _extract_markdown_sections(text: str) -> list[SectionInfo]:
        heading_re = re.compile(r"^(#{1,6})\s+(.+)$", re.MULTILINE)
        matches = list(heading_re.finditer(text))
        if not matches:
            return []
        sections: list[SectionInfo] = []
        for i, m in enumerate(matches):
            level = len(m.group(1))
            heading = m.group(2).strip()
            start = m.end()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
            body = text[start:end].strip()
            if body:
                sections.append(SectionInfo(heading=heading, text=body, level=level))
        return sections

    @staticmethod
    def _extract_plaintext_sections(text: str) -> list[SectionInfo]:
        lines = text.split("\n")
        heading_indices: list[tuple[int, str]] = []
        for i, line in enumerate(lines):
            stripped = line.strip()
            if not stripped or len(stripped) > 120:
                continue
            has_no_ending_punct = stripped[-1] not in ".!?;,:'\"" 
            word_count = len(stripped.split())
            if not has_no_ending_punct or word_count > 15:
                continue

            prev_blank = (i == 0) or (not lines[i - 1].strip())
            next_blank = (i + 1 >= len(lines)) or (not lines[i + 1].strip())
            next_line = lines[i + 1].strip() if i + 1 < len(lines) else ""
            next_much_longer = len(next_line) > len(stripped) * 1.5

            if prev_blank and next_blank:
                heading_indices.append((i, stripped))
            elif word_count <= 10 and next_much_longer:
                heading_indices.append((i, stripped))

        if len(heading_indices) < 2:
            return []

        sections: list[SectionInfo] = []
        for idx, (line_idx, heading) in enumerate(heading_indices):
            start = line_idx + 1
            end = heading_indices[idx + 1][0] if idx + 1 < len(heading_indices) else len(lines)
            body = "\n".join(lines[start:end]).strip()
            if body:
                sections.append(SectionInfo(heading=heading, text=body, level=1))
        return sections
